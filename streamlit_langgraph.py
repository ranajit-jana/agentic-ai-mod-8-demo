import os
import uuid
import operator
import streamlit as st
from typing import TypedDict, Annotated
from langgraph.graph import StateGraph, START, END
from langgraph.checkpoint.memory import MemorySaver
from langgraph.types import interrupt, Command
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from io import BytesIO
from docx import Document
import csv
import io
from fpdf import FPDF
from dotenv import load_dotenv

load_dotenv()

st.set_page_config(page_title="Gemini Research Agent (LangGraph)", layout="centered")

MODEL = "gemini-2.5-flash"
SYSTEM = "You are a helpful research assistant. Provide thorough, well-structured responses."

# LangChain-wrapped Gemini — required for LangGraph streaming (stream_mode="messages")
llm = ChatGoogleGenerativeAI(model=MODEL, google_api_key=os.getenv("GEMINI_API_KEY"))


# --- LangGraph State ---
# TypedDict defines what the graph carries between nodes.
# chat_history uses Annotated[list, operator.add] so each node's returned list
# is APPENDED to the existing history rather than replacing it (LangGraph reducer pattern).

class AgentState(TypedDict):
    user_input: str                              # current question from the human
    outline: str                                 # plan generated for this question
    feedback: str | None                         # human's redirect note (None = approved)
    full_answer: str                             # final generated response
    chat_history: Annotated[list, operator.add]  # accumulated Q&A across all turns


# --- Graph Nodes ---

def plan_node(state: AgentState) -> dict:
    # Generate a short 2-3 bullet outline before answering.
    # Uses chat_history so the plan is aware of prior conversation context.
    messages = [SystemMessage(content=SYSTEM)] + state.get("chat_history", []) + [
        HumanMessage(content=(
            f"The user asked: {state['user_input']}\n\n"
            "Briefly outline in 2-3 bullet points how you will approach answering this. "
            "Plan only — not the full answer yet."
        ))
    ]
    response = llm.invoke(messages)
    # Reset transient fields so a fresh question doesn't carry over stale state
    return {"outline": response.content, "feedback": None, "full_answer": ""}


def human_review_node(state: AgentState) -> dict:
    # interrupt() is the LangGraph HITL primitive.
    # It pauses the graph HERE and saves state to the checkpointer.
    # The value passed to interrupt() is surfaced to the Streamlit UI.
    # When the human decides, graph.invoke(Command(resume=value), config) is called,
    # execution resumes from this exact line with `decision` = value.
    decision = interrupt({"outline": state["outline"]})

    if decision.get("action") == "approve":
        return {"feedback": None}   # route to generate
    else:
        return {"feedback": decision.get("feedback", "")}  # route to revise


def revise_plan_node(state: AgentState) -> dict:
    # Human redirected — revise the outline using their feedback
    messages = [SystemMessage(content=SYSTEM)] + state.get("chat_history", []) + [
        HumanMessage(content=(
            f"The user asked: {state['user_input']}\n\n"
            f"Your previous outline was rejected. Feedback: {state['feedback']}\n\n"
            "Provide a revised 2-3 bullet point outline. Plan only."
        ))
    ]
    response = llm.invoke(messages)
    return {"outline": response.content, "feedback": None}  # clear feedback after revision


def generate_node(state: AgentState) -> dict:
    # Human approved the plan — generate the full answer.
    # LangGraph's stream_mode="messages" will surface tokens from this llm.invoke()
    # call one by one to the Streamlit frontend.
    messages = [SystemMessage(content=SYSTEM)] + state.get("chat_history", []) + [
        HumanMessage(content=f"{state['user_input']}\n\n[Approved plan: {state['outline']}]")
    ]
    response = llm.invoke(messages)
    # Append this Q&A to chat_history — operator.add reducer accumulates it
    return {
        "full_answer": response.content,
        "chat_history": [
            HumanMessage(content=state["user_input"]),
            AIMessage(content=response.content),
        ],
    }


def route_after_review(state: AgentState) -> str:
    # Conditional edge: feedback set → revise the plan; None → generate full answer
    return "revise" if state.get("feedback") else "generate"


# --- Build the Graph ---

builder = StateGraph(AgentState)
builder.add_node("plan", plan_node)
builder.add_node("human_review", human_review_node)
builder.add_node("revise", revise_plan_node)
builder.add_node("generate", generate_node)

builder.add_edge(START, "plan")
builder.add_edge("plan", "human_review")
builder.add_conditional_edges("human_review", route_after_review, {
    "revise": "revise",
    "generate": "generate",
})
builder.add_edge("revise", "human_review")  # loop back so human can review the revised plan
builder.add_edge("generate", END)

def compile_graph():
    # MemorySaver holds the interrupt checkpoint in memory.
    # CRITICAL: must NOT be created at module level — Streamlit reruns the entire
    # script on every interaction, which would create a fresh MemorySaver each time
    # and wipe the interrupted state. Store in st.session_state instead.
    return builder.compile(checkpointer=MemorySaver())


# --- Helpers ---

def get_config(thread_id: str) -> dict:
    # Every graph call must carry the thread_id so LangGraph loads the right checkpoint
    return {"configurable": {"thread_id": thread_id}}


def get_interrupt_value(thread_id: str) -> dict | None:
    # Check whether the graph is currently paused at an interrupt() call.
    # state.next is non-empty when there are pending nodes (i.e. graph was interrupted).
    state = graph.get_state(get_config(thread_id))
    if state.next:
        for task in state.tasks:
            for intr in task.interrupts:
                return intr.value  # {"outline": "..."}
    return None


def stream_generate(thread_id: str, resume_value: dict):
    # Resume the graph after human approval and yield tokens as they arrive.
    # stream_mode="messages" intercepts the LLM call inside generate_node and
    # surfaces each token as a (chunk, metadata) tuple.
    for chunk, metadata in graph.stream(
        Command(resume=resume_value),
        get_config(thread_id),
        stream_mode="messages",
    ):
        # Filter to tokens from the "generate" node only — skip plan/revise noise
        if metadata.get("langgraph_node") == "generate" and getattr(chunk, "content", None):
            yield chunk.content


def session_to_langchain() -> list:
    # Convert st.session_state.messages to LangChain message objects
    # used by quick actions (Subtopics / Summarise) that bypass the graph
    result = []
    for m in st.session_state.messages:
        if m["role"] == "user":
            result.append(HumanMessage(content=m["content"]))
        else:
            result.append(AIMessage(content=m["content"]))
    return result


def stream_direct(prompt: str):
    # Quick actions bypass the graph entirely — call the LLM directly
    messages = [SystemMessage(content=SYSTEM)] + session_to_langchain() + [
        HumanMessage(content=prompt)
    ]
    for chunk in llm.stream(messages):
        if chunk.content:
            yield chunk.content


def reset_conversation():
    st.session_state.messages = []
    # New thread_id = fresh checkpointer state = clean graph
    st.session_state.thread_id = str(uuid.uuid4())


# --- Session state ---
for key, default in {
    "messages": [],
    "thread_id": str(uuid.uuid4()),  # ties this browser session to a LangGraph checkpoint
    "graph": None,
}.items():
    if key not in st.session_state:
        st.session_state[key] = default

# Compile once per session — graph and its MemorySaver live in session_state
if st.session_state.graph is None:
    st.session_state.graph = compile_graph()

graph = st.session_state.graph


# --- UI ---

st.title("Gemini Research Agent")
st.caption("HITL via LangGraph interrupt pattern")

has_messages = bool(st.session_state.messages)
col1, col2, col3 = st.columns(3)

with col1:
    if st.button("Generate Subtopics", use_container_width=True, disabled=not has_messages):
        prompt = "Based on our conversation, generate a list of subtopics in bullet points."
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("assistant"):
            response = st.write_stream(stream_direct(prompt))
        st.session_state.messages.append({"role": "assistant", "content": response})
        st.rerun()

with col2:
    if st.button("Summarise", use_container_width=True, disabled=not has_messages):
        prompt = "Please summarise our conversation so far concisely."
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("assistant"):
            response = st.write_stream(stream_direct(prompt))
        st.session_state.messages.append({"role": "assistant", "content": response})
        st.rerun()

with col3:
    if st.button("Reset", use_container_width=True, type="secondary"):
        reset_conversation()
        st.rerun()

st.divider()

# Replay conversation history
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])


# --- HITL gate (LangGraph interrupt pattern) ---
# get_interrupt_value() reads the checkpointer to see if the graph is paused.
# If it is, we show the plan + Proceed/Redirect UI instead of the chat input.

interrupt_val = get_interrupt_value(st.session_state.thread_id)

if interrupt_val:
    # Graph is paused at human_review_node — show the outline and wait for human decision
    with st.chat_message("assistant"):
        st.markdown("**Here's my plan before I answer:**")
        st.markdown(interrupt_val["outline"])

    st.info("Approve the plan or redirect me.")
    proceed_col, redirect_col = st.columns([1, 2])

    with proceed_col:
        if st.button("Proceed", type="primary", use_container_width=True):
            # Resume the graph with approval — tokens stream from generate_node
            graph_state = graph.get_state(get_config(st.session_state.thread_id))
            user_input = graph_state.values.get("user_input", "")
            st.session_state.messages.append({"role": "user", "content": user_input})
            with st.chat_message("assistant"):
                full_response = st.write_stream(
                    stream_generate(st.session_state.thread_id, {"action": "approve"})
                )
            st.session_state.messages.append({"role": "assistant", "content": full_response})
            st.rerun()

    with redirect_col:
        with st.form("redirect_form"):
            feedback = st.text_input(
                "Redirect:", placeholder="e.g. focus more on X, avoid Y, be more concise..."
            )
            if st.form_submit_button("Update Plan", use_container_width=True):
                if feedback.strip():
                    with st.spinner("Revising plan..."):
                        # Resume with feedback — graph runs revise_plan and interrupts again
                        # graph.invoke() returns as soon as the next interrupt() is hit
                        graph.invoke(
                            Command(resume={"action": "redirect", "feedback": feedback}),
                            get_config(st.session_state.thread_id),
                        )
                    # get_interrupt_value() will now return the revised outline on rerun
                    st.rerun()

else:
    # No active interrupt — graph is idle, accept next question
    if user_input := st.chat_input("Ask a question or continue the conversation..."):
        with st.spinner("Planning..."):
            # Invoke the graph — it runs plan_node then pauses at human_review_node.
            # graph.invoke() returns immediately when interrupt() is hit (doesn't block).
            # Passing chat_history=[] is safe: operator.add accumulates, never clears.
            graph.invoke(
                {
                    "user_input": user_input,
                    "outline": "",
                    "feedback": None,
                    "full_answer": "",
                    "chat_history": [],  # reducer appends; existing history is preserved
                },
                get_config(st.session_state.thread_id),
            )
        # Rerun — get_interrupt_value() will detect the interrupt and show the plan
        st.rerun()


# --- Downloads ---
if st.session_state.messages:
    st.divider()
    st.markdown("### Download Conversation")
    dl1, dl2, dl3 = st.columns(3)

    with dl1:
        doc = Document()
        for m in st.session_state.messages:
            doc.add_heading("You" if m["role"] == "user" else "Assistant", level=2)
            doc.add_paragraph(m["content"])
        word_io = BytesIO()
        doc.save(word_io)
        word_io.seek(0)
        st.download_button(
            "Download Word", data=word_io, file_name="conversation.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    with dl2:
        csv_str_io = io.StringIO()
        writer = csv.writer(csv_str_io)
        writer.writerow(["Role", "Content"])
        for m in st.session_state.messages:
            writer.writerow([m["role"], m["content"]])
        st.download_button(
            "Download CSV", data=BytesIO(csv_str_io.getvalue().encode("utf-8")),
            file_name="conversation.csv", mime="text/csv", use_container_width=True,
        )

    with dl3:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        for m in st.session_state.messages:
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, "You:" if m["role"] == "user" else "Assistant:", ln=True)
            pdf.set_font("Arial", size=11)
            safe = m["content"].encode("latin-1", errors="replace").decode("latin-1")
            pdf.multi_cell(0, 8, safe)
            pdf.ln(4)
        pdf_bytes = pdf.output(dest="S").encode("latin-1")
        st.download_button(
            "Download PDF", data=BytesIO(pdf_bytes),
            file_name="conversation.pdf", mime="application/pdf",
            use_container_width=True,
        )
