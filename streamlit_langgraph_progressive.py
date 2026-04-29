import os
import re
import uuid
import operator
import streamlit as st
from typing import TypedDict, Annotated
from langgraph.graph import StateGraph, START, END
from langgraph.checkpoint.memory import MemorySaver
from langgraph.types import interrupt, Command
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from io import BytesIO
from docx import Document
import csv
import io
from fpdf import FPDF
from dotenv import load_dotenv

load_dotenv()

st.set_page_config(page_title="Progressive Research Agent", layout="wide")

MODEL = "gemini-2.5-flash"
SYSTEM = "You are a helpful research assistant. Provide thorough, well-structured responses."

llm = ChatGoogleGenerativeAI(model=MODEL, google_api_key=os.getenv("GEMINI_API_KEY"))


# ---------------------------------------------------------------------------
# State
# ---------------------------------------------------------------------------
# The graph carries this dict between nodes.
# `sections` is a plain dict (replaced whole each update).
# `chat_history` uses operator.add so returned lists are appended, not replaced.

class AgentState(TypedDict):
    user_input: str                              # research topic
    outline: list                                # list of section title strings
    current_idx: int                             # which section is being worked on
    current_expansion: str                       # just-generated content for current section
    sections: dict                               # {idx: content} — accepted sections
    action: str | None                           # last human decision
    feedback: str | None                         # human's guidance text
    final_answer: str                            # compiled result
    chat_history: Annotated[list, operator.add]  # Q&A history (accumulates across turns)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def parse_outline(text: str) -> list:
    # Strip bullets/numbers and return a clean list of section titles
    result = []
    for line in text.strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        line = re.sub(r'^[\d]+[.)]\s*|^[-•*]\s*', '', line).strip()
        if line:
            result.append(line)
    return result or [text.strip()]


# ---------------------------------------------------------------------------
# Graph nodes
# ---------------------------------------------------------------------------

def plan_node(state: AgentState) -> dict:
    # Generate 3-5 section titles for the research topic.
    # Resets all transient fields so stale state from a previous run doesn't leak.
    messages = [SystemMessage(content=SYSTEM)] + state.get("chat_history", []) + [
        HumanMessage(content=(
            f"Topic: {state['user_input']}\n\n"
            "Generate a structured outline of 3-5 key sections to cover this topic. "
            "Return ONLY the section titles as a numbered list, one per line. "
            "No descriptions, no sub-bullets — just the titles."
        ))
    ]
    response = llm.invoke(messages)
    return {
        "outline": parse_outline(response.content),
        "current_idx": 0,
        "current_expansion": "",
        "sections": {},
        "action": None,
        "feedback": None,
        "final_answer": "",
    }


def outline_review_node(state: AgentState) -> dict:
    # HITL checkpoint 1 — human sees the full outline before any content is written.
    # interrupt() pauses here, saves state to checkpointer, and returns to Streamlit.
    # Revision is handled inline in the UI (update_state), so only "approve" resumes here.
    interrupt({"stage": "outline_review", "outline": state["outline"]})
    return {"action": "approve", "feedback": None}


def expand_section_node(state: AgentState) -> dict:
    # Generate content for the current section.
    # Uses prior approved sections as context so the writing stays coherent.
    # If the human asked to go deeper or rewrite, their feedback is included.
    section_title = state["outline"][state["current_idx"]]

    prior = "\n\n".join(
        f"## {state['outline'][i]}\n{state['sections'][i]}"
        for i in range(state["current_idx"])
        if i in state.get("sections", {})
    )

    parts = [f"Topic: {state['user_input']}", f"Section to expand: {section_title}"]
    if prior:
        parts.append(f"Already written:\n{prior}")
    if state.get("feedback") and state.get("action") in ("deeper", "redirect"):
        parts.append(f"Guidance: {state['feedback']}")
    parts.append(
        "Write a detailed, well-structured expansion of this section (3-5 paragraphs). "
        "Be specific and include examples where relevant."
    )

    messages = [SystemMessage(content=SYSTEM)] + state.get("chat_history", []) + [
        HumanMessage(content="\n\n".join(parts))
    ]
    response = llm.invoke(messages)
    # Clear action/feedback so they don't bleed into the next section
    return {"current_expansion": response.content, "action": None, "feedback": None}


def section_review_node(state: AgentState) -> dict:
    # HITL checkpoint 2 — human reviews each section as it's generated.
    # Runs once per section. Human can approve, skip, go deeper, or rewrite.
    decision = interrupt({
        "stage": "section_review",
        "section_title": state["outline"][state["current_idx"]],
        "section_content": state["current_expansion"],
        "current_idx": state["current_idx"],
        "total": len(state["outline"]),
        "sections": state.get("sections", {}),
        "outline": state["outline"],
    })
    return {
        "action": decision["action"],
        "feedback": decision.get("feedback", None),
    }


def advance_section_node(state: AgentState) -> dict:
    # Accept or skip the current section, then move to the next index.
    sections = dict(state.get("sections", {}))
    if state["action"] == "approve":
        sections[state["current_idx"]] = state["current_expansion"]
    return {
        "sections": sections,
        "current_idx": state["current_idx"] + 1,
        "current_expansion": "",
        "action": None,
        "feedback": None,
    }


def compile_node(state: AgentState) -> dict:
    # Assemble all accepted sections into the final research document.
    # Skipped sections are simply omitted.
    parts = [
        f"## {state['outline'][i]}\n\n{state['sections'][i]}"
        for i in range(len(state["outline"]))
        if i in state.get("sections", {})
    ]
    final_answer = "\n\n---\n\n".join(parts)
    return {
        "final_answer": final_answer,
        "chat_history": [
            HumanMessage(content=state["user_input"]),
            AIMessage(content=final_answer),
        ],
    }


# ---------------------------------------------------------------------------
# Routing
# ---------------------------------------------------------------------------


def route_after_section_review(state: AgentState) -> str:
    # approve/skip → move to next section; deeper/redirect → re-expand same section
    if state["action"] in ("approve", "skip"):
        return "advance_section"
    return "expand_section"


def route_after_advance(state: AgentState) -> str:
    # All sections processed → compile; otherwise expand the next one
    if state["current_idx"] >= len(state["outline"]):
        return "compile"
    return "expand_section"


# ---------------------------------------------------------------------------
# Build graph
# ---------------------------------------------------------------------------

builder = StateGraph(AgentState)
builder.add_node("plan",            plan_node)
builder.add_node("outline_review",  outline_review_node)
builder.add_node("expand_section",  expand_section_node)
builder.add_node("section_review",  section_review_node)
builder.add_node("advance_section", advance_section_node)
builder.add_node("compile",         compile_node)

builder.add_edge(START, "plan")
builder.add_edge("plan", "outline_review")
builder.add_edge("outline_review", "expand_section")  # revision handled inline via update_state
builder.add_edge("expand_section", "section_review")
builder.add_conditional_edges("section_review", route_after_section_review, {
    "advance_section": "advance_section",
    "expand_section":  "expand_section",   # re-expand same section
})
builder.add_conditional_edges("advance_section", route_after_advance, {
    "compile":        "compile",
    "expand_section": "expand_section",
})
builder.add_edge("compile", END)


def compile_graph():
    # MemorySaver MUST be stored in st.session_state — not at module level.
    # Streamlit reruns the entire script on every interaction, so a module-level
    # MemorySaver would be recreated fresh each time, wiping the interrupt state.
    return builder.compile(checkpointer=MemorySaver())


# ---------------------------------------------------------------------------
# Session state
# ---------------------------------------------------------------------------

for key, default in {
    "messages": [],
    "graph":     None,
    "thread_id": str(uuid.uuid4()),
}.items():
    if key not in st.session_state:
        st.session_state[key] = default

if st.session_state.graph is None:
    st.session_state.graph = compile_graph()

graph = st.session_state.graph


# ---------------------------------------------------------------------------
# Graph status helpers
# ---------------------------------------------------------------------------

def get_config(thread_id: str) -> dict:
    return {"configurable": {"thread_id": thread_id}}


def get_graph_status(thread_id: str) -> dict:
    # Reads the checkpointer to determine what stage the graph is at.
    # Uses state.next (which node runs next) instead of task.interrupts
    # because task.interrupts API varies across LangGraph versions.
    try:
        state = graph.get_state(get_config(thread_id))
    except Exception:
        return {"stage": "idle"}

    if not state or not state.values:
        return {"stage": "idle"}

    if state.next:
        v = state.values
        if "outline_review" in state.next:
            return {"stage": "outline_review", "outline": v.get("outline", [])}

        if "section_review" in state.next:
            idx = v.get("current_idx", 0)
            outline = v.get("outline", [])
            return {
                "stage":         "section_review",
                "section_title": outline[idx] if idx < len(outline) else "",
                "section_content": v.get("current_expansion", ""),
                "current_idx":   idx,
                "total":         len(outline),
                "sections":      v.get("sections", {}),
                "outline":       outline,
            }

    if state.values.get("final_answer"):
        return {"stage": "done", "final_answer": state.values["final_answer"]}

    return {"stage": "idle"}


def stream_section(thread_id: str, resume_value: dict):
    # Resume the graph and stream tokens from expand_section_node.
    # stream_mode="messages" intercepts the LLM call token by token.
    # The stream ends when the graph hits the next interrupt() or END.
    for chunk, metadata in graph.stream(
        Command(resume=resume_value),
        get_config(thread_id),
        stream_mode="messages",
    ):
        if metadata.get("langgraph_node") == "expand_section" and getattr(chunk, "content", None):
            yield chunk.content


def reset_conversation():
    st.session_state.messages = []
    st.session_state.thread_id = str(uuid.uuid4())  # fresh thread = fresh checkpointer state


# ---------------------------------------------------------------------------
# UI
# ---------------------------------------------------------------------------

st.title("Progressive Research Agent")
st.caption("LangGraph interrupt — human steers depth and direction at every section")

status = get_graph_status(st.session_state.thread_id)


# ── Stage: IDLE ─────────────────────────────────────────────────────────────
if status["stage"] == "idle":
    if st.session_state.messages:
        # Previous research completed — show it
        for msg in st.session_state.messages:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
        if st.button("Start New Research", type="secondary"):
            reset_conversation()
            st.rerun()
    else:
        st.info(
            "Enter a research topic. The agent generates an outline first — "
            "you approve or redirect it, then review each section one at a time."
        )

    if topic := st.chat_input("What do you want to research?"):
        with st.spinner("Generating outline..."):
            # graph.invoke() returns as soon as outline_review_node hits interrupt()
            graph.invoke(
                {
                    "user_input":        topic,
                    "outline":           [],
                    "current_idx":       0,
                    "current_expansion": "",
                    "sections":          {},
                    "action":            None,
                    "feedback":          None,
                    "final_answer":      "",
                    "chat_history":      [],  # operator.add — passing [] never clears history
                },
                get_config(st.session_state.thread_id),
            )
        st.rerun()


# ── Stage: OUTLINE REVIEW ───────────────────────────────────────────────────
elif status["stage"] == "outline_review":
    st.subheader("Step 1 — Review the Outline")
    st.markdown(
        "The agent will expand each section one at a time. "
        "You can steer the depth and direction at every checkpoint."
    )

    st.markdown("**Proposed sections:**")
    for i, title in enumerate(status["outline"], 1):
        st.markdown(f"{i}. {title}")

    st.divider()
    approve_col, redirect_col = st.columns([1, 2])

    with approve_col:
        if st.button("Approve & Start Writing", type="primary", use_container_width=True):
            # Resume → revise_outline routes to expand_section → stream section 1
            with st.chat_message("assistant"):
                st.write_stream(
                    stream_section(st.session_state.thread_id, {"action": "approve"})
                )
            st.rerun()

    with redirect_col:
        with st.form("outline_redirect"):
            feedback = st.text_area(
                "Redirect the outline:", height=100,
                placeholder=(
                    "e.g. add a section on ethics, remove history, "
                    "focus more on practical applications..."
                )
            )
            if st.form_submit_button("Revise Outline", use_container_width=True):
                if feedback.strip():
                    with st.spinner("Revising..."):
                        # Generate the revised outline directly via LLM and
                        # patch it into the graph state via update_state().
                        # This keeps the graph interrupted at outline_review_node
                        # (same interrupt, updated outline value) — more reliable
                        # than resuming into a second sequential interrupt().
                        old_outline = status["outline"]
                        user_input = graph.get_state(
                            get_config(st.session_state.thread_id)
                        ).values.get("user_input", "")
                        revision_prompt = (
                            f"Topic: {user_input}\n\n"
                            "Previous outline:\n"
                            + "\n".join(f"{i+1}. {s}" for i, s in enumerate(old_outline))
                            + f"\n\nFeedback: {feedback}\n\n"
                            "Revise the outline based on the feedback. "
                            "Return ONLY numbered section titles, one per line."
                        )
                        response = llm.invoke([
                            SystemMessage(content=SYSTEM),
                            HumanMessage(content=revision_prompt),
                        ])
                        new_outline = parse_outline(response.content)
                        # Patch the outline in state — graph stays interrupted at outline_review
                        graph.update_state(
                            get_config(st.session_state.thread_id),
                            {"outline": new_outline},
                        )
                st.rerun()


# ── Stage: SECTION REVIEW ───────────────────────────────────────────────────
elif status["stage"] == "section_review":
    current_idx = status["current_idx"]
    total       = status["total"]

    # Progress
    st.progress(current_idx / total, text=f"Section {current_idx + 1} of {total}")

    # Approved sections so far
    if status["sections"]:
        with st.expander(
            f"Approved so far — {len(status['sections'])} of {total} sections", expanded=False
        ):
            for i, title in enumerate(status["outline"]):
                if i in status["sections"]:
                    st.markdown(f"**✓ {title}**")
                    st.markdown(status["sections"][i])
                    if i < len(status["outline"]) - 1:
                        st.divider()

    # Current section
    st.subheader(f"Section {current_idx + 1}: {status['section_title']}")
    with st.chat_message("assistant"):
        st.markdown(status["section_content"])

    st.divider()
    st.markdown("**What would you like to do with this section?**")

    left_col, right_col = st.columns([1, 2])

    with left_col:
        # Approve — save section, stream next (or compile if last)
        if st.button("Approve", type="primary", use_container_width=True):
            label = "Expanding next section..." if current_idx + 1 < total else "Compiling..."
            with st.spinner(label):
                with st.chat_message("assistant"):
                    st.write_stream(
                        stream_section(st.session_state.thread_id, {"action": "approve"})
                    )
            st.rerun()

        # Skip — discard section, stream next (or compile if last)
        if st.button("Skip this section", use_container_width=True):
            label = "Expanding next section..." if current_idx + 1 < total else "Compiling..."
            with st.spinner(label):
                with st.chat_message("assistant"):
                    st.write_stream(
                        stream_section(st.session_state.thread_id, {"action": "skip"})
                    )
            st.rerun()

    with right_col:
        with st.form("section_action_form"):
            feedback = st.text_area(
                "Optional guidance:", height=90,
                placeholder=(
                    "e.g. go deeper on X, add examples, "
                    "focus on Y, be more concise..."
                )
            )
            go_deeper_col, rewrite_col = st.columns(2)

            with go_deeper_col:
                go_deeper = st.form_submit_button("Go Deeper", use_container_width=True)
            with rewrite_col:
                rewrite = st.form_submit_button("Rewrite", use_container_width=True)

            if go_deeper:
                guidance = feedback.strip() or "Expand with significantly more depth, detail, and concrete examples."
                with st.spinner("Going deeper..."):
                    with st.chat_message("assistant"):
                        st.write_stream(
                            stream_section(
                                st.session_state.thread_id,
                                {"action": "deeper", "feedback": guidance},
                            )
                        )
                st.rerun()

            if rewrite:
                if feedback.strip():
                    with st.spinner("Rewriting..."):
                        with st.chat_message("assistant"):
                            st.write_stream(
                                stream_section(
                                    st.session_state.thread_id,
                                    {"action": "redirect", "feedback": feedback},
                                )
                            )
                    st.rerun()
                else:
                    st.warning("Enter guidance to rewrite the section.")

    if st.button("Abandon & Start Over", type="secondary"):
        reset_conversation()
        st.rerun()


# ── Stage: DONE ─────────────────────────────────────────────────────────────
elif status["stage"] == "done":
    st.success("Research complete!")

    with st.chat_message("assistant"):
        st.markdown(status["final_answer"])

    # Store in messages so it's shown on next idle render
    st.session_state.messages = [{"role": "assistant", "content": status["final_answer"]}]

    st.divider()
    st.markdown("### Download")
    dl1, dl2, dl3 = st.columns(3)
    content = status["final_answer"]

    with dl1:
        doc = Document()
        for line in content.split("\n"):
            if line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.strip() == "---":
                doc.add_paragraph("")
            else:
                doc.add_paragraph(line)
        word_io = BytesIO()
        doc.save(word_io)
        word_io.seek(0)
        st.download_button(
            "Download Word", data=word_io, file_name="research.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    with dl2:
        csv_str_io = io.StringIO()
        writer = csv.writer(csv_str_io)
        writer.writerow(["Section", "Content"])
        for chunk in content.split("\n---\n"):
            lines = chunk.strip().split("\n", 1)
            title = lines[0].replace("## ", "") if lines else ""
            body  = lines[1].strip() if len(lines) > 1 else ""
            writer.writerow([title, body])
        st.download_button(
            "Download CSV",
            data=BytesIO(csv_str_io.getvalue().encode("utf-8")),
            file_name="research.csv", mime="text/csv", use_container_width=True,
        )

    with dl3:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        for line in content.split("\n"):
            if line.startswith("## "):
                pdf.set_font("Arial", "B", 13)
                safe = line[3:].encode("latin-1", errors="replace").decode("latin-1")
                pdf.cell(0, 10, safe, ln=True)
            elif line.strip() == "---":
                pdf.ln(4)
            else:
                pdf.set_font("Arial", size=11)
                safe = line.encode("latin-1", errors="replace").decode("latin-1")
                pdf.multi_cell(0, 7, safe)
        pdf_bytes = pdf.output(dest="S").encode("latin-1")
        st.download_button(
            "Download PDF", data=BytesIO(pdf_bytes),
            file_name="research.pdf", mime="application/pdf",
            use_container_width=True,
        )

    st.divider()
    if st.button("Start New Research", type="secondary"):
        reset_conversation()
        st.rerun()
